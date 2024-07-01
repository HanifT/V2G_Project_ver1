from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Mathematical Formulation of the Pyomo Model', level=1)

# Nomenclature section
doc.add_heading('Nomenclature', level=2)
doc.add_heading('Sets:', level=3)
sets = [
    "T: Set of time periods.",
    "V: Set of vehicles."
]
for item in sets:
    doc.add_paragraph(item)

doc.add_heading('Parameters:', level=3)
parameters = [
    "CRTP_t: Cost of electricity at time period t.",
    "GHG_t: Greenhouse gas emission rate at time period t.",
    "trv_dist_{v,t}: Travel distance for vehicle v at time period t.",
    "SOC_REQ_{v,t}: State of charge requirement for vehicle v at time period t.",
    "fac_chr_{v,t}: Charging indicator for vehicle v at time period t.",
    "lev_chr_{v,t}: Charge type for vehicle v at time period t.",
    "veh_model_{v,t}: Vehicle model for vehicle v at time period t.",
    "location_{v,t}: Location of vehicle v at time period t.",
    "bat_cap_{v,t}: Battery capacity of vehicle v at time period t.",
    "MAX_{v,t}: Maximum charging rate for vehicle v at time period t.",
    "C_THRESHOLD: Minimum SOC buffer during charging/discharging.",
    "eff_chr: Charging efficiency.",
    "eff_dri_{v,t}: Driving efficiency for vehicle v at time period t based on the vehicle model.",
    "ghg_cost: Cost per gram of greenhouse gas emissions.",
    "d_slope_{v,t}: Slope of battery degradation for vehicle v at time period t.",
    "d_intercept_{v,t}: Intercept of battery degradation for vehicle v at time period t."
]
for item in parameters:
    doc.add_paragraph(item)

doc.add_heading('Decision Variables:', level=3)
decision_variables = [
    "X_CHR_{v,t}: Charging power for vehicle v at time period t.",
    "SOC_{v,t}: State of charge for vehicle v at time period t.",
    "batt_deg_{v,t}: Battery degradation for vehicle v at time period t.",
    "cumulative_charging_{v,t}: Cumulative charging for vehicle v at time period t.",
    "X_CHR_neg_part_{v,t}: Negative part of X_CHR_{v,t} for vehicle v at time period t.",
    "batt_deg_cost_{v,t}: Battery degradation cost for vehicle v at time period t."
]
for item in decision_variables:
    doc.add_paragraph(item)

# Objective function section
doc.add_heading('Objective Function:', level=2)
objective_function = """
Minimize: 
    sum_{v in V} sum_{t in T} (CRTP_t * X_CHR_{v,t} / 1000 + GHG_t * X_CHR_{v,t} * ghg_cost / 1000 + batt_deg_cost_{v,t})
"""
doc.add_paragraph(objective_function)

# Constraints section
doc.add_heading('Constraints:', level=2)

constraints = [
    {
        "title": "State of Charge Balance:",
        "formula": """
SOC_{v,t} = 
    if t = 0:
        100
    else:
        SOC_{v,t-1} + (X_CHR_{v,t} * eff_chr / bat_cap_{v,t} - trv_dist_{v,t} / (bat_cap_{v,t} * eff_dri_{v,t})) * 100
"""
    },
    {
        "title": "Minimum Charging Rate:",
        "formula": """
X_CHR_{v,t} >= 
    if location_{v,t} in locs:
        -MAX_{v,t}
    else:
        0
"""
    },
    {
        "title": "Maximum Charging Rate:",
        "formula": """
X_CHR_{v,t} <= MAX_{v,t} * fac_chr_{v,t}
"""
    },
    {
        "title": "Minimum SOC at Departure:",
        "formula": """
SOC_{v,t} >= SOC_REQ_{v,t} + 20 if t = 0 or fac_chr_{v,t-1} = 1 and fac_chr_{v,t} = 0
"""
    },
    {
        "title": "SOC Buffer:",
        "formula": """
SOC_{v,t} >= C_THRESHOLD if fac_chr_{v,t} = 1
"""
    },
    {
        "title": "Zero Charging when Indicator is Zero:",
        "formula": """
X_CHR_{v,t} = 0 if fac_chr_{v,t} = 0
"""
    },
    {
        "title": "Capture Negative Part of Charging:",
        "formula": """
X_CHR_neg_part_{v,t} >= X_CHR_{v,t} if fac_chr_{v,t} = 1
"""
    },
    {
        "title": "Cumulative Charging:",
        "formula": """
cumulative_charging_{v,t} = 
    if t = 0:
        0
    else:
        cumulative_charging_{v,t-1} + X_CHR_neg_part_{v,t} * fac_chr_{v,t}
"""
    },
    {
        "title": "Battery Degradation:",
        "formula": """
batt_deg_{v,t} = 
    if fac_chr_{v,t} = 1:
        d_slope_{v,t} * cumulative_charging_{v,t}
    else if fac_chr_{v,t} = 0:
        0
"""
    },
    {
        "title": "Battery Degradation Cost:",
        "formula": """
batt_deg_cost_{v,t} = 
    if fac_chr_{v,t} = 1:
        batt_deg_{v,t} - batt_deg_{v,t-1}
    else if fac_chr_{v,t} = 0:
        0
"""
    }
]

for constraint in constraints:
    doc.add_heading(constraint["title"], level=3)
    doc.add_paragraph(constraint["formula"])

# Save the document
doc.save("Pyomo_Model_Formulation.docx")


